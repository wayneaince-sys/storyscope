"""Generate a 'rough draft' (v1) and a 'revision' (v2) manuscript for Compare/PDF QA.

v1 is intentionally abstraction-heavy and repetitive so suggestions fire.
v2 trims, rewrites a few paragraphs concretely, and reorders one chapter.
"""
from docx import Document

def write(path, chapters):
    doc = Document()
    for title, paragraphs in chapters:
        doc.add_heading(title, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
    doc.save(path)


# v1 — rough draft, abstract & repetitive
V1 = [
    ("Chapter 1: The Letter", [
        # repetitive abstract opener
        "Margaret felt a strange sense of dread as she contemplated the meaning of the letter. She felt fear and confusion and an overwhelming feeling that her life would change. She thought about the implications. She felt the weight of memory and history and regret pressing down on her in the silence of the hallway.",
        # echo of paragraph 1 — same sentiment-laden abstractions
        "Margaret felt a deep sense of dread when she considered what the letter meant. She felt afraid and confused and overwhelmed by a feeling that everything was changing. She thought about the implications and felt the weight of memory pressing on her.",
        "She thought about the past and the future and the meaning of family. She thought about her father and the unspoken regrets between them. She thought about how everything was about to be different.",
        "The clock ticked. The radiator hummed. She did not move.",
    ]),
    ("Chapter 2: A Small Funeral", [
        # abstract-heavy
        "Grief was a quiet, abstract weight in Margaret's chest. She felt the strangeness of mourning a man she had not understood. She thought about regret, about absence, about the meaning of family obligation, about whether love and duty were the same thing.",
        # low-tension flat
        "The cemetery was a place. The priest spoke. People stood. Margaret thought about other things while the words went on. The sky was a sky. The wind moved a little. After some time the service ended and the people began to leave the cemetery in an orderly fashion.",
        "Afterward a woman touched her elbow and said something kind about her father, and Margaret nodded politely without really hearing what was said.",
    ]),
    ("Chapter 3: The Workshop", [
        "The workshop was full of memory and history and the lingering presence of her father's labor. Margaret felt the weight of his absence. She thought about the meaning of inheritance, about what it meant to receive a person's tools after their death.",
        # echo of chapter 1 paragraph
        "She felt a strange sense of dread again as she contemplated the meaning of being there. She felt fear and confusion and a feeling that her life was changing. She thought about implications. She felt memory and history and regret pressing down.",
        "Inside a wooden box she found a key.",
    ]),
    ("Chapter 4: The House on Lark Street", [
        "Lark Street was narrow and wet with rain. Margaret walked slowly, counting house numbers. The address was 47, a tall thin building with peeling green paint and a brass knocker shaped like a fox.",
        "She pushed the key into the lock. It turned with a soft click. The hallway inside was dark and smelled of cedar.",
        "She thought she heard movement on the floor above. She stilled her breath and listened. The house held its silence like something alive.",
    ]),
    ("Chapter 5: What Was Found", [
        # abstract / told
        "On the second floor she discovered the truth about her father's life. It was a complicated truth, full of ambiguity and contradiction. She experienced a profound shift in her understanding of who he had been and what he had done.",
        # low-tension flat
        "She read the papers. There were many papers. Reading them took a long time. The information they contained was significant. She processed the information slowly, methodically, without rushing.",
        "She closed the last folder and looked up. The rain had stopped. She felt a strange sense of dread, but also of release.",
    ]),
]

# v2 — revision: tightened, more concrete, reordered, with new chapter 6
V2 = [
    ("Chapter 1: The Letter", [
        # tighter, concrete
        "Margaret found the envelope on the porch step, half-buried under the dead leaves. The paper was thick, yellowed at the edges, and there was no return address — only her name in slanted black ink.",
        # kept short, sensory
        "Her cold hands turned it over once, then again. The kitchen clock ticked. The radiator hummed. She could not bring herself to open it.",
        "She stood in the hallway for a long moment, the unopened letter pressed flat against her thigh, and listened to the house breathe around her.",
    ]),
    ("Chapter 2: A Small Funeral", [
        # concrete, sensory
        "The cemetery was small and damp. Margaret stood at the back, behind a row of strangers in black coats. The priest's voice carried thin in the wind. A crow shifted on a branch, then lifted away.",
        "Her last conversation with her father had been about the weather. She kept thinking about that — the small size of it, the politeness, the way she had said goodbye without meaning anything by it.",
        "After the service, a woman touched her elbow. 'I knew him,' the woman said softly. She had pale grey eyes and held a bundle wrapped in brown paper. 'He wanted you to have this.'",
    ]),
    ("Chapter 3: The Workshop", [
        "The workshop smelled of varnish and sawdust. A single yellow lamp swung from the rafters. Margaret stepped over a coil of rope and let her fingers trail across the workbench. The wood was warm.",
        "She unwrapped the bundle on the bench. Inside was a small wooden box, lacquered black, with a brass clasp. The box, when she opened it, was lined with red velvet. Nestled in the velvet was a single brass key.",
        # kept one repeated 'strange weight' for shared-repetition signal
        "She felt a strange weight in her chest, an old dread she could not place.",
    ]),
    ("Chapter 4: The House on Lark Street", [
        "Lark Street was narrow and wet with rain. Margaret walked slowly, counting house numbers. Number 47 was a tall thin building with peeling green paint and a brass knocker shaped like a fox.",
        "The key turned with a soft click. The hallway inside was dark and smelled of cedar. A staircase climbed up into shadow.",
        "She thought she heard movement on the floor above. She stilled her breath and listened. The house held its silence like something alive.",
    ]),
    ("Chapter 5: What Was Found", [
        "On the second floor she found a room of papers — letters, photographs, ledgers, a typewriter under a dusty cloth. Her father's handwriting filled the margins of everything. He had been keeping a record. He had been waiting for her.",
        "She sat on the floor and read until her hands were cold and her eyes ached. The story he had been writing was about her mother. About the night her mother had gone, and what he had done in the days after.",
        "She closed the last folder and looked up. The rain had stopped. The light through the window was thin and silver. She felt a strange weight in her chest, an old dread she could not place — though now she knew exactly what it was.",
    ]),
    # NEW chapter in v2
    ("Chapter 6: The Return", [
        "She walked home with the typewriter in her arms, ink-smudged ledgers tucked under one elbow. The streetlamps had come on. Her own breath made a small white shape in front of her face, then disappeared.",
        "On the porch step she paused, set the typewriter down on the boards, and looked back at Lark Street. The fox knocker glinted faintly under the rain. She lifted the key in her palm, weighed it once, and put it in her pocket.",
        "Inside, she set the kettle on. The kitchen clock ticked. The radiator hummed. This time, she opened the letter.",
    ]),
]

write("/home/user/workspace/storyscope/sample/sample_v1.docx", V1)
write("/home/user/workspace/storyscope/sample/sample_v2.docx", V2)
print("Wrote sample_v1.docx and sample_v2.docx")
