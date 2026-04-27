"""Generate a small sample manuscript for QA."""
from docx import Document

doc = Document()

CHAPTERS = [
    ("Chapter 1: The Letter", [
        "Margaret found the envelope on the porch step, half-buried under the dead leaves. The paper was thick and yellowed at the edges. She hesitated, then turned it over in her cold hands. There was no return address, only her name in slanted black ink.",
        "She felt a strange weight in her chest, an old dread she could not place. She knew, in some deep part of herself, that this letter was going to change everything. She took it inside, closed the door, and stood in the hallway for a long moment, unable to move.",
        "The kitchen clock ticked. The radiator hummed. She could not bring herself to open it.",
    ]),
    ("Chapter 2: A Small Funeral", [
        "The cemetery was small and damp. Margaret stood at the back, behind a row of strangers. The priest's voice carried thin in the wind. She watched a crow shift on a branch, then lift away.",
        "She thought about how little she had known her father. Their last conversation had been about the weather. She wondered what it meant to mourn someone you barely understood.",
        "After the service, a woman touched her elbow. 'I knew him,' the woman said softly. Margaret turned. The woman had pale grey eyes and held a bundle wrapped in brown paper. 'He wanted you to have this.'",
    ]),
    ("Chapter 3: The Workshop", [
        "The workshop smelled of varnish and sawdust. A single yellow lamp swung from the rafters. Margaret stepped over a coil of rope and let her fingers trail across the workbench. The wood was warm.",
        "She unwrapped the bundle carefully. Inside was a small wooden box, lacquered black, with a brass clasp. She set it on the bench and listened to the quiet building around her.",
        "She felt a strange weight in her chest, an old dread she could not place. The box, when she opened it, was lined with red velvet. Nestled in the velvet was a key.",
    ]),
    ("Chapter 4: The House on Lark Street", [
        "Lark Street was narrow and wet with rain. Margaret walked slowly, counting house numbers. The address her father had left her was 47, a tall thin building with peeling green paint and a brass knocker shaped like a fox.",
        "She pushed the key into the lock. It turned with a soft click. The hallway inside was dark and smelled of cedar. A staircase climbed up into shadow.",
        "She thought she heard movement on the floor above. She stilled her breath and listened. The house held its silence like something alive.",
    ]),
    ("Chapter 5: What Was Found", [
        "On the second floor she found a room of papers — letters, photographs, ledgers, a typewriter under a dusty cloth. Her father's handwriting filled the margins of everything. He had been keeping a record. He had been waiting for her.",
        "She sat on the floor and read until her hands were cold and her eyes ached. The story he had been writing was about her mother. About the night her mother had gone, and what he had done in the days after.",
        "She closed the last folder and looked up. The rain had stopped. The light through the window was thin and silver. She felt a strange weight in her chest, an old dread she could not place — though now she knew exactly what it was.",
    ]),
]

for title, paragraphs in CHAPTERS:
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)

doc.save("sample.docx")
print("Wrote sample.docx")
